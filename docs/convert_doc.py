#!/usr/bin/env python3
"""Convert FaultRay system document Markdown to formatted Word document."""

import re
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

INPUT_MD = Path(__file__).parent / "FaultRay_System_Document.md"
OUTPUT_DOCX = Path(__file__).parent / "FaultRay_System_Document.docx"

# Theme colors
NAVY = RGBColor(0x1A, 0x3C, 0x6E)
DARK_BLUE = RGBColor(0x2C, 0x5F, 0x9E)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
DARK_GRAY = RGBColor(0x33, 0x33, 0x33)
CODE_BG = "F0F4F8"
FONT_NAME = "Yu Gothic"
FONT_NAME_CODE = "Consolas"
FONT_SIZE = Pt(10)


def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_run_font(run, bold=False, size=None, color=None, font_name=FONT_NAME, italic=False):
    run.font.name = font_name
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="{font_name}"/>')
        rPr.insert(0, rFonts)
    else:
        rFonts.set(qn("w:eastAsia"), font_name)
    if bold:
        run.bold = True
    if italic:
        run.italic = True
    if size:
        run.font.size = size
    if color:
        run.font.color.rgb = color


def add_heading_styled(doc, text, level):
    p = doc.add_heading(text, level=level)
    sizes = {1: Pt(18), 2: Pt(14), 3: Pt(12), 4: Pt(11)}
    for run in p.runs:
        set_run_font(run, bold=True, color=NAVY, size=sizes.get(level, Pt(11)))
    return p


def add_paragraph_with_bold(doc, text, style=None):
    p = doc.add_paragraph(style=style)
    # Handle **bold** and `code` markers
    parts = re.split(r"(\*\*[^*]+\*\*|`[^`]+`)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = p.add_run(part[2:-2])
            set_run_font(run, bold=True, size=FONT_SIZE, color=DARK_GRAY)
        elif part.startswith("`") and part.endswith("`"):
            run = p.add_run(part[1:-1])
            set_run_font(run, size=Pt(9), font_name=FONT_NAME_CODE, color=DARK_BLUE)
        else:
            run = p.add_run(part)
            set_run_font(run, size=FONT_SIZE, color=DARK_GRAY)
    return p


def add_code_block(doc, lines):
    """Add a code block with monospace font and light background."""
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.left_indent = Cm(1)
        run = p.add_run(line)
        set_run_font(run, size=Pt(8.5), font_name=FONT_NAME_CODE, color=RGBColor(0x2D, 0x3A, 0x4A))


def add_table_from_rows(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(header)
        set_run_font(run, bold=True, size=Pt(9), color=WHITE)
        set_cell_shading(cell, "1A3C6E")

    for row_idx, row_data in enumerate(rows):
        bg = "FFFFFF" if row_idx % 2 == 0 else "F5F7FA"
        for col_idx in range(min(len(row_data), len(headers))):
            cell_text = row_data[col_idx]
            cell = table.rows[row_idx + 1].cells[col_idx]
            cell.text = ""
            set_cell_shading(cell, bg)
            # Handle **bold** and `code`
            parts = re.split(r"(\*\*[^*]+\*\*|`[^`]+`)", cell_text)
            for part in parts:
                if part.startswith("**") and part.endswith("**"):
                    run = cell.paragraphs[0].add_run(part[2:-2])
                    set_run_font(run, bold=True, size=Pt(9))
                elif part.startswith("`") and part.endswith("`"):
                    run = cell.paragraphs[0].add_run(part[1:-1])
                    set_run_font(run, size=Pt(8), font_name=FONT_NAME_CODE, color=DARK_BLUE)
                else:
                    run = cell.paragraphs[0].add_run(part)
                    set_run_font(run, size=Pt(9))

    return table


def parse_md_table(lines, start_idx):
    headers = [h.strip() for h in lines[start_idx].strip().strip("|").split("|")]
    rows = []
    idx = start_idx + 2  # Skip separator line
    while idx < len(lines) and lines[idx].strip().startswith("|"):
        row = [c.strip() for c in lines[idx].strip().strip("|").split("|")]
        rows.append(row)
        idx += 1
    return headers, rows, idx


def convert():
    content = INPUT_MD.read_text(encoding="utf-8")
    lines = content.split("\n")
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # Default font
    style = doc.styles["Normal"]
    style.font.name = FONT_NAME
    style.font.size = FONT_SIZE
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="{FONT_NAME}"/>')
        rPr.insert(0, rFonts)
    else:
        rFonts.set(qn("w:eastAsia"), FONT_NAME)

    i = 0
    in_code_block = False
    code_lines = []

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Code blocks
        if stripped.startswith("```"):
            if in_code_block:
                add_code_block(doc, code_lines)
                code_lines = []
                in_code_block = False
            else:
                in_code_block = True
            i += 1
            continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        # Empty lines
        if not stripped:
            i += 1
            continue

        # Horizontal rule
        if stripped == "---":
            # Add a thin line
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            i += 1
            continue

        # Headings
        if stripped.startswith("# ") and not stripped.startswith("## "):
            add_heading_styled(doc, stripped[2:], 1)
            i += 1
            continue

        if stripped.startswith("## "):
            add_heading_styled(doc, stripped[3:], 2)
            i += 1
            continue

        if stripped.startswith("### "):
            add_heading_styled(doc, stripped[4:], 3)
            i += 1
            continue

        if stripped.startswith("#### "):
            p = doc.add_paragraph()
            run = p.add_run(stripped[5:])
            set_run_font(run, bold=True, size=Pt(10), color=NAVY)
            i += 1
            continue

        # Tables
        if stripped.startswith("|") and i + 1 < len(lines) and "---" in lines[i + 1]:
            headers, rows, end_idx = parse_md_table(lines, i)
            add_table_from_rows(doc, headers, rows)
            doc.add_paragraph()
            i = end_idx
            continue

        # Bullet lists
        if stripped.startswith("- ") or stripped.startswith("* "):
            text = stripped[2:]
            add_paragraph_with_bold(doc, text, style="List Bullet")
            i += 1
            continue

        # Numbered lists
        if re.match(r"^\d+\.\s", stripped):
            text = re.sub(r"^\d+\.\s", "", stripped)
            add_paragraph_with_bold(doc, text, style="List Number")
            i += 1
            continue

        # Regular paragraph
        add_paragraph_with_bold(doc, stripped)
        i += 1

    doc.save(str(OUTPUT_DOCX))
    print(f"Saved: {OUTPUT_DOCX}")


if __name__ == "__main__":
    convert()
